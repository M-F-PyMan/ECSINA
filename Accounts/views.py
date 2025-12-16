from rest_framework import viewsets, status, permissions
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.exceptions import PermissionDenied
from django.utils.dateparse import parse_date
from django.utils import timezone
from datetime import timedelta
import logging
from django.core.cache import cache
from django.contrib.auth import authenticate
from suggestion.utils.cache_helpers import invalidate_related_educations_cache
from django.http import FileResponse
from django.db.models import Q
from rest_framework.decorators import action
from docx import Document
from django.http import HttpResponse
from xhtml2pdf import pisa
from rest_framework_simplejwt.tokens import RefreshToken
from docx.oxml.ns import qn
from django.conf import settings
from docx.shared import Pt
from pathlib import Path
from django.core.validators import validate_email
from django.core.exceptions import ValidationError
from .models import (
    User,
    Profile,
    TempUser,
    UserProposalUpload,
    ConsultantReply,
    UserCategory,
    UserCategoryItem,
    UserActivityLog,
    OTP

)

from .serializers import (
    UserSerializer,
    ProfileSerializer,
    TempUserSerializer,
    UserProposalUploadSerializer,
    ConsultantReplySerializer,
    UserCategorySerializer,
    UserCategoryItemSerializer,
    UserActivityLogSerializer,
    RegisterSerializer,
    SendOTPSerializer,
    VerifyOTPSerializer,
    ResetPasswordSerializer,

)
from django_ratelimit.decorators import ratelimit
from permissions.permissions import IsAdminOrOwner
import logging

logger = logging.getLogger(__name__)

# ========= Existing viewsets kept as-is (snipped for brevity); ensure they remain in the file =========
# UserActivityLogViewSet, UserViewSet, ProfileViewSet, TempUserViewSet, DeleteUserView,
# UserProposalUploadViewSet, ConsultantReplyViewSet, UserCategoryViewSet, UserCategoryItemViewSet
# RegisterView, SendOTPView, VerifyOTPView, ResetPasswordView
# ================================================================================================

# --- Security parameters for login lockout ---
LOGIN_MAX_ATTEMPTS = getattr(settings, "LOGIN_MAX_ATTEMPTS", 5)
LOGIN_LOCK_MINUTES = getattr(settings, "LOGIN_LOCK_MINUTES", 15)

class UserActivityLogViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = UserActivityLog.objects.select_related('user', 'file').order_by('-timestamp')
    serializer_class = UserActivityLogSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return self.queryset.filter(user=self.request.user)


class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all().order_by('-created_at')
    serializer_class = UserSerializer
    permission_classes = [IsAdminOrOwner]

    def get_queryset(self):
        queryset = super().get_queryset()
        role = self.request.query_params.get('role')
        email = self.request.query_params.get('email')
        is_active = self.request.query_params.get('active')
        from_date = self.request.query_params.get('from_date')
        to_date = self.request.query_params.get('to_date')

        if role:
            queryset = queryset.filter(role__name=role)
        if email:
            queryset = queryset.filter(email__icontains=email)
        if is_active in ['true', 'false']:
            queryset = queryset.filter(is_active=(is_active == 'true'))
        if from_date:
            queryset = queryset.filter(created_at__date__gte=parse_date(from_date))
        if to_date:
            queryset = queryset.filter(created_at__date__lte=parse_date(to_date))

        return queryset

    def update(self, request, *args, **kwargs):
        user = request.user
        data = request.data.copy()
        if not user.role or user.role.name not in ['owner', 'admin']:
            data.pop('role_id', None)
            data.pop('is_active', None)
            data.pop('is_staff', None)
        request._full_data = data
        return super().update(request, *args, **kwargs)

    def partial_update(self, request, *args, **kwargs):
        return self.update(request, *args, **kwargs)


class ProfileViewSet(viewsets.ModelViewSet):
    queryset = Profile.objects.select_related('user').all().order_by('-user__created_at')
    serializer_class = ProfileSerializer
    permission_classes = [IsAdminOrOwner]


class TempUserViewSet(viewsets.ModelViewSet):
    queryset = TempUser.objects.all().order_by('-sms_sent_date')
    serializer_class = TempUserSerializer
    permission_classes = [IsAdminOrOwner]


class DeleteUserView(APIView):
    permission_classes = [IsAdminOrOwner]

    def delete(self, request, user_id):
        try:
            user_to_delete = User.objects.get(id=user_id)
            if user_to_delete == request.user:
                return Response({"detail": "امکان حذف خودتان وجود ندارد."}, status=status.HTTP_400_BAD_REQUEST)
            user_to_delete.deleted_by = request.user
            user_to_delete.save()
            user_to_delete.delete()
            return Response({"detail": f"کاربر با موفقیت حذف شد توسط {request.user.name}."}, status=status.HTTP_204_NO_CONTENT)
        except User.DoesNotExist:
            return Response({"detail": "کاربر پیدا نشد."}, status=status.HTTP_404_NOT_FOUND)


class UserProposalUploadViewSet(viewsets.ModelViewSet):
    queryset = UserProposalUpload.objects.select_related('user', 'original_proposal').all().order_by('-uploaded_at')
    serializer_class = UserProposalUploadSerializer
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]

    def get_queryset(self):
        user = self.request.user
        queryset = super().get_queryset()

        status_param = self.request.query_params.get('status')
        name = self.request.query_params.get('name')
        email = self.request.query_params.get('email')
        from_date = self.request.query_params.get('from_date')
        to_date = self.request.query_params.get('to_date')

        if status_param:
            queryset = queryset.filter(status=status_param)
        if name:
            queryset = queryset.filter(user__name__icontains=name)
        if email:
            queryset = queryset.filter(user__email__icontains=email)
        if from_date:
            queryset = queryset.filter(uploaded_at__date__gte=parse_date(from_date))
        if to_date:
            queryset = queryset.filter(uploaded_at__date__lte=parse_date(to_date))

        if user.role and user.role.name in ['admin', 'owner']:
            return queryset
        return queryset.filter(Q(user=user) | Q(is_free=True))

    def perform_create(self, serializer):
        instance = serializer.save(user=self.request.user)
        category_items = UserCategoryItem.objects.filter(file=instance)
        for item in category_items:
            invalidate_related_educations_cache(item.user_category.user.id, item.user_category.id)

    def perform_update(self, serializer):
        instance = self.get_object()
        if not instance.can_user_edit(self.request.user):
            raise PermissionDenied("شما اجازه ویرایش این فایل را ندارید.")
        updated = serializer.save()
        category_items = UserCategoryItem.objects.filter(file=updated)
        for item in category_items:
            invalidate_related_educations_cache(item.user_category.user.id, item.user_category.id)

    @action(detail=True, methods=['get'], url_path='download')
    def download_file(self, request, pk=None):
        instance = self.get_object()
        if not instance.can_user_download(request.user):
            raise PermissionDenied("شما اجازه دانلود این فایل را ندارید.")
        return FileResponse(instance.file.open(), as_attachment=True, filename=instance.file.name)

    @action(detail=True, methods=['patch'], url_path='save-edit')
    def save_editable_content(self, request, pk=None):
        instance = self.get_object()
        if not instance.can_user_edit(request.user):
            raise PermissionDenied("شما اجازه ویرایش این فایل را ندارید.")
        html = request.data.get('editable_content')
        instance.editable_content = html
        instance.save()
        return Response({'status': 'saved'})

    @action(detail=True, methods=['get'], url_path='export-word')
    def export_word(self, request, pk=None):
        instance = self.get_object()
        if not instance.can_user_download(request.user):
            raise PermissionDenied("شما اجازه دریافت این فایل را ندارید.")

        doc = Document()
        run = doc.add_paragraph().add_run(instance.editable_content or "بدون محتوا")

        # استخراج نام فونت از مسیر
        font_name = Path(instance.font_path).stem if instance.font_path else 'Vazir'
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(14)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{instance.title}.docx"'
        doc.save(response)
        return response

    @action(detail=True, methods=['get'], url_path='export-pdf')
    def export_pdf(self, request, pk=None):
        instance = self.get_object()

        if not instance.can_user_download(request.user):
            raise PermissionDenied("شما اجازه دریافت این فایل را ندارید.")

        html_content = instance.editable_content or "<p>بدون محتوا</p>"

        # اضافه کردن فونت انتخاب‌شده به CSS
        font_path = instance.font_path or '/static/fonts/vazir.ttf'
        font_css = f"""
           <style>
           @font-face {{
               font-family: CustomFont;
               src: url('{font_path}');
           }}
           body {{
               font-family: CustomFont;
               direction: rtl;
           }}
           </style>
           """

        full_html = f"""
           <html>
           <head>{font_css}</head>
           <body>{html_content}</body>
           </html>
           """

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{instance.title}.pdf"'
        pisa.CreatePDF(full_html, dest=response)
        return response

class ConsultantReplyViewSet(viewsets.ModelViewSet):
    queryset = ConsultantReply.objects.select_related('consultant', 'proposal_upload__user').all().order_by('-replied_at')
    serializer_class = ConsultantReplySerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        queryset = super().get_queryset()

        from_date = self.request.query_params.get('from_date')
        to_date = self.request.query_params.get('to_date')

        if from_date:
            queryset = queryset.filter(replied_at__date__gte=parse_date(from_date))
        if to_date:
            queryset = queryset.filter(replied_at__date__lte=parse_date(to_date))

        # حالت استثنایی: هم مشاور هم ادمین
        if getattr(user, 'is_consultant_admin', False):
            return queryset

        if user.role and user.role.name == 'consultant':
            return queryset.filter(consultant=user)

        if user.role and user.role.name in ['admin', 'owner']:
            return queryset

        return queryset.filter(
            Q(proposal_upload__user=user) |
            Q(proposal_upload__is_free=True)
        )

    def perform_create(self, serializer):
        user = self.request.user
        if not user.role or user.role.name != 'consultant':
            raise PermissionDenied("فقط مشاورها می‌تونن پاسخ بدن.")
        serializer.save(consultant=user)


class UserCategoryViewSet(viewsets.ModelViewSet):
    serializer_class = UserCategorySerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return UserCategory.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)


class UserCategoryItemViewSet(viewsets.ModelViewSet):
    serializer_class = UserCategoryItemSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return UserCategoryItem.objects.filter(user_category__user=self.request.user)

    def perform_create(self, serializer):
        user_category = serializer.validated_data['user_category']
        file = serializer.validated_data['file']
        if user_category.user != self.request.user:
            raise PermissionDenied("شما مجاز به افزودن آیتم به این دسته‌بندی نیستید.")
        if UserCategoryItem.objects.filter(user_category=user_category, file=file).exists():
            raise PermissionDenied("این فایل قبلاً در این دسته‌بندی اضافه شده.")
        instance = serializer.save()
        invalidate_related_educations_cache(self.request.user.id, instance.user_category_id)

    def perform_destroy(self, instance):
        invalidate_related_educations_cache(self.request.user.id, instance.user_category_id)
        instance.delete()



class RegisterView(APIView):
    def post(self, request):
        serializer = RegisterSerializer(data=request.data)
        if serializer.is_valid():
            user = serializer.save()
            refresh = RefreshToken.for_user(user)
            return Response({
                'message': 'ثبت‌نام با موفقیت انجام شد',
                'access': str(refresh.access_token),
                'refresh': str(refresh)
            }, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class SendOTPView(APIView):
    def post(self, request):
        serializer = SendOTPSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        identifier = serializer.validated_data['identifier']

        user = User.objects.get(email=identifier) if "@" in identifier else User.objects.get(mobile=identifier)
        otp = OTP.create_for_user(user)
        result = otp.send_sms(api_key=settings.FARAZSMS_API_KEY)

        if not result.get("success"):
            return Response({"detail": result.get("error")}, status=status.HTTP_502_BAD_GATEWAY)

        return Response({"detail": "کد تأیید ارسال شد."}, status=status.HTTP_200_OK)


class VerifyOTPView(APIView):
    def post(self, request):
        serializer = VerifyOTPSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        identifier = serializer.validated_data['identifier']
        code = serializer.validated_data['code']

        user = User.objects.get(email=identifier) if "@" in identifier else User.objects.get(mobile=identifier)
        otp = OTP.objects.filter(user=user, code=code).order_by('-created_at').first()

        if not otp or not otp.is_valid():
            user.sms_sent_tries += 1
            user.save(update_fields=['sms_sent_tries'])
            return Response({"detail": "کد نامعتبر یا منقضی شده."}, status=status.HTTP_400_BAD_REQUEST)

        otp.mark_used()
        return Response({"detail": "کد تأیید شد. می‌توانید رمز را تغییر دهید."}, status=status.HTTP_200_OK)


class ResetPasswordView(APIView):
    def post(self, request):
        serializer = ResetPasswordSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        identifier = serializer.validated_data['identifier']
        new_password = serializer.validated_data['new_password']

        user = User.objects.get(email=identifier) if "@" in identifier else User.objects.get(mobile=identifier)
        user.set_password(new_password)
        user.save()

        return Response({"detail": "رمز عبور با موفقیت تغییر کرد."}, status=status.HTTP_200_OK)





class LoginView(APIView):
    @ratelimit(key='ip', rate='10/m', block=True)
    def post(self, request):
        identifier = request.data.get('username')  # email or mobile
        password = request.data.get('password')

        if not identifier or not password:
            return Response({'detail': 'نام کاربری یا رمز عبور نامعتبر است'}, status=status.HTTP_400_BAD_REQUEST)

        ip = request.META.get('REMOTE_ADDR', '')
        cache_key = f"login_fail:{identifier}:{ip}"
        fail = cache.get(cache_key) or {'count': 0, 'locked_until': None}

        # Check lock
        if fail['locked_until'] and timezone.now() < fail['locked_until']:
            return Response({'detail': 'حساب موقتاً قفل شده است. بعداً تلاش کنید.'}, status=status.HTTP_429_TOO_MANY_REQUESTS)

        # Find user (email or mobile)
        try:
            user = User.objects.get(email=identifier) if "@" in identifier else User.objects.get(mobile=identifier)
        except User.DoesNotExist:
            # Increment fail count and maybe lock
            fail['count'] += 1
            if fail['count'] >= LOGIN_MAX_ATTEMPTS:
                fail['locked_until'] = timezone.now() + timedelta(minutes=LOGIN_LOCK_MINUTES)
            cache.set(cache_key, fail, timeout=LOGIN_LOCK_MINUTES * 60)
            logger.warning("Failed login: identifier not found. ip=%s", ip)
            return Response({'detail': 'نام کاربری یا رمز عبور نامعتبر است'}, status=status.HTTP_401_UNAUTHORIZED)

        # Check password and active status
        if not user.check_password(password) or not user.is_active:
            fail['count'] += 1
            if fail['count'] >= LOGIN_MAX_ATTEMPTS:
                fail['locked_until'] = timezone.now() + timedelta(minutes=LOGIN_LOCK_MINUTES)
            cache.set(cache_key, fail, timeout=LOGIN_LOCK_MINUTES * 60)
            logger.warning("Failed login: bad password or inactive. user_id=%s ip=%s", user.id, ip)
            return Response({'detail': 'نام کاربری یا رمز عبور نامعتبر است'}, status=status.HTTP_401_UNAUTHORIZED)

        # Success: clear fail counter
        cache.delete(cache_key)

        refresh = RefreshToken.for_user(user)
        access_token = str(refresh.access_token)
        refresh_token = str(refresh)

        # Set HttpOnly, Secure cookie for refresh
        resp = Response({
            'access': access_token,
            'user_id': user.id,
            'email': user.email,
            'mobile': user.mobile,
        }, status=status.HTTP_200_OK)

        cookie_max_age = 60 * 60 * 24 * 7  # 7 days
        resp.set_cookie(
            key='refresh_token',
            value=refresh_token,
            max_age=cookie_max_age,
            httponly=True,
            secure=not settings.DEBUG,
            samesite='Lax',
            path='/api/'
        )

        logger.info("User logged in. user_id=%s ip=%s", user.id, ip)
        return resp


class RefreshTokenView(APIView):
    def post(self, request):
        refresh_token = request.COOKIES.get('refresh_token')
        if not refresh_token:
            return Response({'detail': 'Refresh token missing'}, status=status.HTTP_401_UNAUTHORIZED)
        try:
            token = RefreshToken(refresh_token)
            access = str(token.access_token)
            return Response({'access': access}, status=status.HTTP_200_OK)
        except Exception:
            return Response({'detail': 'Invalid refresh token'}, status=status.HTTP_401_UNAUTHORIZED)


class LogoutView(APIView):
    def post(self, request):
        resp = Response({'detail': 'logged out'}, status=status.HTTP_200_OK)
        # Match cookie path used in set_cookie
        resp.delete_cookie('refresh_token', path='/api/')
        return resp
